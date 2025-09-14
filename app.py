from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
import requests
import io
import json
from io import BytesIO
import os
import fitz
import re

from copy import deepcopy
import zipfile
from urllib.parse import urlparse
from lxml import etree
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


app = Flask(__name__)

# Allow all origins for now - can restrict later to specific domains
CORS(app, resources={r"/*": {"origins": "*"}})

# -----------------------
# TEST ENDPOINT
# -----------------------
@app.route("/", methods=["GET"])
def home():
    return jsonify({"status": "API is running", "message": "Welcome to LegallyAI API!"}), 200


# -----------------------
# Utility Functions
# -----------------------

def load_numbering_map(file_bytes):
    """Parses numbering.xml and builds a map: numId -> list_type (bullet/ordered)."""
    numbering_map = {}
    with zipfile.ZipFile(file_bytes) as z:
        if "word/numbering.xml" not in z.namelist():
            return numbering_map  # no lists in docx

        xml = z.read("word/numbering.xml")
        root = etree.fromstring(xml)

        for num in root.findall(".//w:num", namespaces=root.nsmap):
            num_id = num.get(qn("w:numId"))
            if not num_id:
                continue
            abstract_elem = num.find("w:abstractNumId", namespaces=root.nsmap)
            if abstract_elem is None:
                continue
            abstract_id = abstract_elem.get(qn("w:val"))

            # find abstract definition
            abstract = root.find(f".//w:abstractNum[@w:abstractNumId='{abstract_id}']", namespaces=root.nsmap)
            if abstract is not None:
                fmt = abstract.find(".//w:numFmt", namespaces=root.nsmap)
                if fmt is not None:
                    fmt_val = fmt.get(qn("w:val"))
                    if fmt_val == "bullet":
                        numbering_map[int(num_id)] = "bullet"
                    else:
                        numbering_map[int(num_id)] = "ordered"
    return numbering_map


# -----------------------------
# STEP 1: Build numId -> format map
# -----------------------------
def build_num_format_map(doc):
    try:
        numbering_part = doc.part.numbering_part.element
    except KeyError:
        return {}

    tree = etree.fromstring(etree.tostring(numbering_part))
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    num_format_map = {}
    for num in tree.findall(".//w:num", namespaces=ns):
        num_id = num.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId")
        abstract_num_id = num.find(".//w:abstractNumId", namespaces=ns).get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")

        abstract_num = tree.find(f".//w:abstractNum[@w:abstractNumId='{abstract_num_id}']", namespaces=ns)
        if abstract_num is None:
            continue

        fmt_elem = abstract_num.find(".//w:numFmt", namespaces=ns)
        fmt = fmt_elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") if fmt_elem is not None else None

        # --- Detect if uppercase via <w:lvl><w:lvlText> ---
        lvl = abstract_num.find(".//w:lvl", namespaces=ns)
        if lvl is not None:
            lvl_text_elem = lvl.find(".//w:lvlText", namespaces=ns)
            if lvl_text_elem is not None:
                lvl_text_val = lvl_text_elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                if fmt == "lowerLetter" and "%1" in lvl_text_val and lvl_text_val.strip().startswith("%1") and lvl_text_val.isupper():
                    fmt = "upperLetter"
                elif fmt == "lowerRoman" and "%1" in lvl_text_val and lvl_text_val.isupper():
                    fmt = "upperRoman"

        num_format_map[num_id] = fmt

    return num_format_map


# -----------------------------
# STEP 2: Detect list type from paragraph
# -----------------------------
def get_list_type_and_indent(paragraph, num_format_map):
    p = paragraph._element
    numPr = p.find(".//w:numPr", namespaces=p.nsmap)
    if numPr is None:
        return None, 0, None

    ilvl = numPr.find(".//w:ilvl", namespaces=p.nsmap)
    numId = numPr.find(".//w:numId", namespaces=p.nsmap)

    indent = int(ilvl.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")) if ilvl is not None else 0
    num_id = numId.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") if numId is not None else None

    fmt = num_format_map.get(num_id, None)

    # --- Deep dive: check the actual <w:lvl> for overrides ---
    if num_id is not None:
        try:
            numbering_part = paragraph.part.numbering_part.element
            tree = etree.fromstring(etree.tostring(numbering_part))
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

            abstract_num_id_elem = tree.find(f".//w:num[@w:numId='{num_id}']/w:abstractNumId", namespaces=ns)
            if abstract_num_id_elem is not None:
                abstract_num_id = abstract_num_id_elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                lvl = tree.find(f".//w:abstractNum[@w:abstractNumId='{abstract_num_id}']/w:lvl[@w:ilvl='{indent}']", namespaces=ns)
                if lvl is not None:
                    # Prefer <w:numFmt> at the level
                    lvl_fmt_elem = lvl.find(".//w:numFmt", namespaces=ns)
                    if lvl_fmt_elem is not None:
                        fmt = lvl_fmt_elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")

                    # Double-check <w:lvlText> for case (uppercase letters/romans)
                    lvl_text_elem = lvl.find(".//w:lvlText", namespaces=ns)
                    if lvl_text_elem is not None:
                        lvl_text_val = lvl_text_elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                        if fmt == "lowerLetter" and lvl_text_val and lvl_text_val.isupper():
                            fmt = "upperLetter"
                        elif fmt == "lowerRoman" and lvl_text_val and lvl_text_val.isupper():
                            fmt = "upperRoman"
        except Exception:
            pass

    # --- Map to Quill-compatible list type ---
    if fmt in ["decimal", "lowerLetter", "upperLetter", "lowerRoman", "upperRoman"]:
        list_type = "ordered"
    elif fmt == "bullet":
        list_type = "bullet"
    else:
        list_type = None

    return list_type, indent, fmt


# -----------------------------
# STEP 3: Extract full docx to Quill Delta
# -----------------------------
def extract_docx_to_delta(file_bytes):
    doc = Document(file_bytes)
    num_format_map = build_num_format_map(doc)  # your existing function

    delta = []

    # --- Process paragraphs ---
    for para in doc.paragraphs:
        # --- Detect alignment ---
        paragraph_alignment = para.alignment
        align_str = None
        if paragraph_alignment == WD_ALIGN_PARAGRAPH.CENTER:
            align_str = "center"
        elif paragraph_alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            align_str = "right"
        elif paragraph_alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            align_str = "justify"

        # --- Detect list type + indent ---
        list_type, indent, fmt = get_list_type_and_indent(para, num_format_map)

        # --- Process runs ---
        for run in para.runs:
            attrs = {}
            if run.bold: attrs["bold"] = True
            if run.italic: attrs["italic"] = True
            if run.underline: attrs["underline"] = True
            if run.font and run.font.strike: attrs["strike"] = True
            if run.font and run.font.superscript: attrs["script"] = "super"
            if run.font and run.font.subscript: attrs["script"] = "sub"
            if run.font and run.font.name: attrs["font"] = run.font.name

            if run.font and run.font.size:
                try:
                    pt = run.font.size.pt
                    if pt <= 10:
                        attrs["size"] = "small"
                    elif pt >= 16:
                        attrs["size"] = "large"
                except:
                    pass

            if run.font and run.font.color and run.font.color.rgb:
                attrs["color"] = f"#{run.font.color.rgb}"

            if run.text.strip():
                insert_obj = {"insert": run.text}
                if attrs:
                    insert_obj["attributes"] = attrs
                delta.append(insert_obj)

        # --- Paragraph break ---
        paragraph_break = {"insert": "\n"}
        attrs = {}
        if align_str:
            attrs["align"] = align_str
        if list_type:
            attrs["list"] = list_type
            if indent > 0:
                attrs["indent"] = indent
            if fmt:
                attrs["numFmt"] = fmt
        if attrs:
            paragraph_break["attributes"] = attrs
        delta.append(paragraph_break)

    # --- Process tables ---
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            row_map = {}
            for i, cell in enumerate(row.cells):
                # Convert each paragraph in the cell to Quill delta
                cell_delta = []
                for para in cell.paragraphs:
                    for run in para.runs:
                        attrs = {}
                        if run.bold: attrs["bold"] = True
                        if run.italic: attrs["italic"] = True
                        if run.underline: attrs["underline"] = True
                        if run.font and run.font.strike: attrs["strike"] = True
                        if run.font and run.font.superscript: attrs["script"] = "super"
                        if run.font and run.font.subscript: attrs["script"] = "sub"
                        if run.font and run.font.name: attrs["font"] = run.font.name
                        if run.font and run.font.size:
                            pt = run.font.size.pt
                            if pt <= 10: attrs["size"] = "small"
                            elif pt >= 16: attrs["size"] = "large"
                        if run.font and run.font.color and run.font.color.rgb:
                            attrs["color"] = f"#{run.font.color.rgb}"

                        if run.text.strip():
                            insert_obj = {"insert": run.text}
                            if attrs:
                                insert_obj["attributes"] = deepcopy(attrs)
                            cell_delta.append(insert_obj)

                    # Paragraph break
                    cell_delta.append({"insert": "\n"})

                row_map[f"col{i}"] = cell_delta
            table_rows.append(row_map)

        delta.append({
            "insert": {
                "table": {
                    "rows": table_rows
                }
            }
        })
        delta.append({"insert": "\n"})


    return delta



# --- Helper: process a single paragraph into delta ops ---
def process_paragraph(para, num_format_map):
    ops = []

    # --- Detect alignment ---
    paragraph_alignment = para.alignment
    align_str = None
    if paragraph_alignment == WD_ALIGN_PARAGRAPH.CENTER:
        align_str = "center"
    elif paragraph_alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        align_str = "right"
    elif paragraph_alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        align_str = "justify"

    # --- Detect list type + indent ---
    list_type, indent, fmt = get_list_type_and_indent(para, num_format_map)

    # --- Process runs ---
    for run in para.runs:
        attrs = {}
        if run.bold: attrs["bold"] = True
        if run.italic: attrs["italic"] = True
        if run.underline: attrs["underline"] = True
        if run.font and run.font.strike: attrs["strike"] = True
        if run.font and run.font.superscript: attrs["script"] = "super"
        if run.font and run.font.subscript: attrs["script"] = "sub"
        if run.font and run.font.name: attrs["font"] = run.font.name

        if run.font and run.font.size:
            try:
                pt = run.font.size.pt
                if pt <= 10:
                    attrs["size"] = "small"
                elif pt >= 16:
                    attrs["size"] = "large"
            except:
                pass

        if run.font and run.font.color and run.font.color.rgb:
            attrs["color"] = f"#{run.font.color.rgb}"

        if run.text.strip():
            insert_obj = {"insert": run.text}
            if attrs:
                insert_obj["attributes"] = attrs
            ops.append(insert_obj)

    # --- Paragraph break ---
    paragraph_break = {"insert": "\n"}
    attrs = {}
    if align_str:
        attrs["align"] = align_str
    if list_type:
        attrs["list"] = list_type
        if indent > 0:
            attrs["indent"] = indent
        if fmt:
            attrs["numFmt"] = fmt
    if attrs:
        paragraph_break["attributes"] = attrs
    ops.append(paragraph_break)

    return ops



def detect_list_type_and_indent(text, x0):
    """
    Detect list type and nesting level from text markers + x position.
    """
    list_type = None
    indent = max(0, int(x0 // 40))  # fallback indent from margin
    skip_chars = 0

    # --- Bullet markers ---
    if text.strip() in ("•", "-", "‣", "▪", "*"):
        list_type = "bullet"
        skip_chars = len(text)
        return list_type, indent, skip_chars

    # --- Multi-level numeric lists (1., 1.1., 1.1.1.) ---
    if re.match(r"^\d+(\.\d+)*[\.\)]\s*", text):
        list_type = "ordered"
        # nesting = number of dots = indent
        indent = text.count(".")
        skip_chars = len(re.match(r"^\d+(\.\d+)*[\.\)]\s*", text).group(0))
        return list_type, indent, skip_chars

    # --- Alphabetical lists: a., (a), A. ---
    if re.match(r"^\(?[a-zA-Z]\)?[\.\)]\s*", text):
        list_type = "ordered"
        indent = 1
        skip_chars = len(re.match(r"^\(?[a-zA-Z]\)?[\.\)]\s*", text).group(0))
        return list_type, indent, skip_chars

    # --- Roman numerals: i., (iv), IV. ---
    if re.match(r"^\(?[ivxlcdmIVXLCDM]+\)?[\.\)]\s*", text):
        list_type = "ordered"
        indent = 1
        skip_chars = len(re.match(r"^\(?[ivxlcdmIVXLCDM]+\)?[\.\)]\s*", text).group(0))
        return list_type, indent, skip_chars

    return None, indent, 0


def extract_pdf_to_delta(file_bytes):
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    delta = []
    margin = 40  # tolerance for alignment detection

    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block["type"] != 0:
                continue

            for line in block["lines"]:
                if not line["spans"]:
                    continue

                first_span_text = line["spans"][0]["text"]
                x0, y0, x1, y1 = line["bbox"]

                # --- Detect list type + indent ---
                list_type, indent, skip_chars = detect_list_type_and_indent(first_span_text, x0)

                # --- Detect alignment ---
                page_width = page.rect.width
                if abs((x0 + x1) / 2 - page_width / 2) < margin:
                    align = "center"
                elif page_width - x1 < margin:
                    align = "right"
                elif x0 < margin:
                    align = "left"
                else:
                    align = "justify"

                # --- Process spans ---
                first_span_processed = False
                for span in line["spans"]:
                    text = span["text"]
                    if not text.strip():
                        continue

                    if not first_span_processed and list_type and skip_chars > 0:
                        text = text[skip_chars:].lstrip()
                        first_span_processed = True

                    if not text:
                        continue

                    attrs = {}
                    font_name = span.get("font", "")
                    if "Bold" in font_name:
                        attrs["bold"] = True
                    if "Italic" in font_name or "Oblique" in font_name:
                        attrs["italic"] = True
                    if font_name:
                        attrs["font"] = font_name

                    size = span.get("size", 12)
                    if size <= 10:
                        attrs["size"] = "small"
                    elif size >= 16:
                        attrs["size"] = "large"

                    color = span.get("color", None)
                    if color:
                        hex_color = "#{:06x}".format(color)
                        attrs["color"] = hex_color

                    insert_obj = {"insert": text}
                    if attrs:
                        insert_obj["attributes"] = attrs
                    delta.append(insert_obj)

                # --- Add paragraph break ---
                paragraph_break = {"insert": "\n"}
                attrs = {"align": align}
                if list_type:
                    attrs["list"] = list_type
                    if indent > 0:
                        attrs["indent"] = indent
                if attrs:
                    paragraph_break["attributes"] = attrs
                delta.append(paragraph_break)

    return delta



@app.route('/extract', methods=['POST'])
def extract():
    try:
        file_url = request.json.get('fileUrl')
        if not file_url:
            return jsonify({"error": "No fileUrl provided"}), 400

        # Download the file
        response = requests.get(file_url)
        if response.status_code != 200:
            return jsonify({"error": "Failed to fetch the file"}), 400

        file_bytes = BytesIO(response.content)

        # Determine file extension
        parsed_url = urlparse(file_url)
        filename = os.path.basename(parsed_url.path)
        ext = os.path.splitext(filename)[1].lower()

        if ext == ".docx":
            delta = extract_docx_to_delta(file_bytes)
        elif ext == ".pdf":
            delta = extract_pdf_to_delta(file_bytes)
        else:
            return jsonify({"error": f"Unsupported file type: {ext}"}), 400

        return jsonify({"delta": delta})

    except Exception as e:
        return jsonify({"error": str(e)}), 500

#if __name__ == '__main__':
    #app.run(debug=True)
    
if __name__ == "__main__":
    app.run(debug=False, use_reloader=False)

